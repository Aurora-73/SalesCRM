r"""Convert .doc/.docx files to Markdown.

Usage:
    python tools/doc2md.py path\to\file.docx
    python tools/doc2md.py path\to\directory
    python tools/doc2md.py

When no path is provided, the legacy default directory is used:
    docs/other-book/doc

On successful conversion this script keeps the existing behavior and removes
the source Word file.
"""

from __future__ import annotations

import glob
import html
import os
import shutil
import stat
import sys
import tempfile
import zipfile
import xml.etree.ElementTree as ET

from docx import Document


PROJECT_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
DEFAULT_DOC_DIR = os.path.join(PROJECT_ROOT, "docs", "other-book", "doc")
WORD_XML_NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}


def extract_docx_xml_text(docx_path: str) -> str:
    """Fallback extractor for text inside Word text boxes and drawing objects."""
    try:
        with zipfile.ZipFile(docx_path) as archive:
            xml_bytes = archive.read("word/document.xml")
    except Exception:
        return ""

    root = ET.fromstring(xml_bytes)
    lines: list[str] = []

    for paragraph in root.iterfind(".//w:p", WORD_XML_NS):
        texts = [
            html.unescape(node.text)
            for node in paragraph.iterfind(".//w:t", WORD_XML_NS)
            if node.text
        ]
        line = "".join(texts).strip()
        if line:
            lines.append(line)

    return "\n".join(lines).strip() + "\n" if lines else ""


def docx_to_markdown(docx_path: str) -> str:
    """Extract simple structured text from .docx and render it as Markdown."""
    doc = Document(docx_path)
    lines: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            lines.append("")
            continue

        style_name = (para.style.name or "").lower() if para.style else ""
        if style_name.startswith("heading"):
            try:
                level = int(style_name.replace("heading", "").strip())
            except ValueError:
                level = 1
            level = min(max(level, 1), 6)
            lines.append(f"{'#' * level} {text}")
        elif style_name.startswith("list"):
            lines.append(f"- {text}")
        else:
            lines.append(text)

    for table in doc.tables:
        if lines and lines[-1] != "":
            lines.append("")
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = " ".join(
                    para.text.strip() for para in cell.paragraphs if para.text.strip()
                )
                cells.append(cell_text)
            row_text = " | ".join(cell for cell in cells if cell)
            if row_text:
                lines.append(row_text)

    content = "\n".join(lines).strip()
    if content:
        return content + "\n"

    return extract_docx_xml_text(docx_path)


def convert_doc_via_word(doc_path: str, tmp_dir: str) -> str | None:
    """Convert a .doc file to a temporary .docx file through Word COM."""
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(doc_path)
        tmp_docx = os.path.join(
            tmp_dir, os.path.splitext(os.path.basename(doc_path))[0] + ".docx"
        )
        doc.SaveAs2(tmp_docx, FileFormat=16)
        doc.Close()
        return tmp_docx
    except Exception as exc:
        print(f"  Word COM convert failed: {exc}")
        return None
    finally:
        word.Quit()


def process_file(filepath: str, tmp_dir: str) -> bool:
    ext = os.path.splitext(filepath)[1].lower()
    md_path = os.path.splitext(filepath)[0] + ".md"

    if ext == ".docx":
        try:
            md_content = docx_to_markdown(filepath)
        except Exception as exc:
            print(f"  python-docx failed: {exc}")
            return False
    elif ext == ".doc":
        tmp_docx = convert_doc_via_word(filepath, tmp_dir)
        if tmp_docx is None:
            return False
        try:
            md_content = docx_to_markdown(tmp_docx)
        except Exception as exc:
            print(f"  Converted .docx parse failed: {exc}")
            return False
    else:
        return False

    if not md_content.strip():
        return False

    with open(md_path, "w", encoding="utf-8", newline="\n") as handle:
        handle.write(md_content)

    os.chmod(filepath, stat.S_IWRITE)
    os.remove(filepath)
    return True


def collect_input_files(paths: list[str]) -> list[str]:
    files: list[str] = []

    for path in paths:
        full_path = os.path.abspath(path)
        if os.path.isdir(full_path):
            files.extend(glob.glob(os.path.join(full_path, "*.doc")))
            files.extend(glob.glob(os.path.join(full_path, "*.docx")))
        elif os.path.isfile(full_path):
            if os.path.splitext(full_path)[1].lower() in (".doc", ".docx"):
                files.append(full_path)
            else:
                print(f"Skip non-Word file: {path}")
        else:
            print(f"Path does not exist: {path}")

    return sorted({f for f in files if not f.lower().endswith(".md")})


def default_files() -> list[str]:
    if not os.path.isdir(DEFAULT_DOC_DIR):
        print(f"Directory does not exist: {DEFAULT_DOC_DIR}")
        sys.exit(1)

    return sorted(
        glob.glob(os.path.join(DEFAULT_DOC_DIR, "*.doc"))
        + glob.glob(os.path.join(DEFAULT_DOC_DIR, "*.docx"))
    )


def main() -> None:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    sys.stderr.reconfigure(encoding="utf-8", errors="replace")

    files = collect_input_files(sys.argv[1:]) if len(sys.argv) > 1 else default_files()

    print(f"Found {len(files)} file(s) to convert\n")

    tmp_dir = tempfile.mkdtemp(prefix="doc2md_")
    ok, fail = 0, 0

    try:
        for index, filepath in enumerate(files, 1):
            name = os.path.basename(filepath)
            print(f"[{index}/{len(files)}] {name} ... ", end="", flush=True)
            try:
                if process_file(filepath, tmp_dir):
                    print("OK")
                    ok += 1
                else:
                    print("FAIL")
                    fail += 1
            except Exception as exc:
                print(f"ERROR: {exc}")
                fail += 1
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)

    print(f"\nDone: {ok} succeeded, {fail} failed")


if __name__ == "__main__":
    main()
